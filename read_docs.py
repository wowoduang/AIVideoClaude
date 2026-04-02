from docx import Document
import os
import sys

# 设置控制台编码为 UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def read_docx(filepath):
    """读取 Word 文档并返回所有段落文本"""
    try:
        doc = Document(filepath)
        content = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                content.append(f"段落 {i+1}: {para.text}")
        
        # 读取表格内容
        for table in doc.tables:
            content.append("\n【表格内容】")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        return f"读取失败：{e}"

if __name__ == "__main__":
    docs = [f for f in os.listdir('.') if f.endswith('.docx')]
    print(f"\n找到 {len(docs)} 个 Word 文档\n")
    
    doc_order = [
        "NarratoAI_PRD_v2.1_剧情块主线修订版.docx",
        "NarratoAI_PRD_v2.1_story_block_revision.docx", 
        "technical_consensus_after_last_doc.docx",
        "剧情分段与视频边界融合规则草稿_v0.1.docx"
    ]
    
    for idx, doc_name in enumerate(doc_order, 1):
        if doc_name in docs:
            print(f"\n{'='*80}")
            print(f"文档 {idx}/4: {doc_name}")
            print('='*80)
            
            content = read_docx(doc_name)
            lines = content.split('\n')
            
            # 智能输出：显示前 60 行关键内容
            keywords = ['修订', '模块', '剧情块', '字幕', '视觉', '表格', '场景', '视频', '音频', '时间轴']
            displayed_count = 0
            
            for line in lines:
                if displayed_count >= 60:
                    break
                # 检查是否包含关键词
                if any(keyword in line for keyword in keywords):
                    print(line)
                    displayed_count += 1
            
            print(f"\n... (共{len(lines)}行内容，已显示{displayed_count}行关键内容)")
        else:
            print(f"\n⚠️ 未找到文档：{doc_name}")
